import os
import shutil
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.db.models import Document
from app.services.historique_service import enregistrer_action

# Racine du projet = 3 niveaux au-dessus de ce fichier
# (document_service.py -> app/services -> app -> racine du projet)
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
UPLOAD_DIR = os.path.join(BASE_DIR, "data", "uploads")


def save_uploaded_file(file_path: str, original_filename: str) -> str:
    """
    Copie le fichier uploadé dans le dossier data/uploads/
    et retourne le chemin final (absolu) où il est stocké.
    Si le fichier est déjà dans uploads/, ne le recopie pas.
    """
    os.makedirs(UPLOAD_DIR, exist_ok=True)
    destination = os.path.join(UPLOAD_DIR, original_filename)

    if os.path.abspath(file_path) != os.path.abspath(destination):
        shutil.copy(file_path, destination)

    return destination


def extract_text_from_pdf(file_path: str) -> str:
    """Extrait tout le texte d'un fichier PDF."""
    reader = PdfReader(file_path)
    texte = ""
    for page in reader.pages:
        contenu = page.extract_text()
        if contenu:
            texte += contenu + "\n"
    return texte.strip()


def extract_text_from_docx(file_path: str) -> str:
    """Extrait tout le texte d'un fichier Word (.docx)."""
    doc = DocxDocument(file_path)
    texte = "\n".join([p.text for p in doc.paragraphs])
    return texte.strip()


def extract_text(file_path: str) -> str:
    """
    Détecte le type de fichier (PDF ou DOCX) et extrait le texte
    avec la bonne méthode.
    """
    extension = file_path.lower().split(".")[-1]

    if extension == "pdf":
        return extract_text_from_pdf(file_path)
    elif extension == "docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Type de fichier non supporté : {extension}")


def create_document_entry(
    db: Session,
    file_path: str,
    original_filename: str,
    uploaded_by: int,
    categorie_id: int = None
) -> Document:
    """
    Processus complet : sauvegarde le fichier, extrait le texte,
    et enregistre le document dans la base de données.
    """
    chemin_final = save_uploaded_file(file_path, original_filename)
    texte = extract_text(chemin_final)

    extension = original_filename.lower().split(".")[-1]
    taille_ko = os.path.getsize(chemin_final) // 1024

    nouveau_document = Document(
        nom_original=original_filename,
        chemin_fichier=chemin_final,
        type_fichier=extension,
        taille=taille_ko,
        categorie_id=categorie_id,
        uploaded_by=uploaded_by,
        contenu_extrait=texte,
    )

    db.add(nouveau_document)
    db.commit()
    db.refresh(nouveau_document)

    # Traçabilité : on enregistre cette action dans l'historique
    enregistrer_action(
        db=db,
        utilisateur_id=uploaded_by,
        action="upload",
        document_id=nouveau_document.id,
        details=f"Dépôt du document '{original_filename}'"
    )

    return nouveau_document


def get_document_path(db: Session, document_id: int, utilisateur_id: int, role: str = "utilisateur") -> str:
    """
    Retourne le chemin du fichier pour permettre son téléchargement,
    et enregistre cette action dans l'historique.
    Un utilisateur non-admin ne peut accéder qu'à ses propres documents.
    """
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise ValueError(f"Document avec l'id {document_id} introuvable.")

    if role != "admin" and document.uploaded_by != utilisateur_id:
        raise PermissionError("Vous n'avez pas accès à ce document.")

    if not os.path.exists(document.chemin_fichier):
        raise FileNotFoundError(f"Le fichier n'existe plus à l'emplacement : {document.chemin_fichier}")

    enregistrer_action(
        db=db,
        utilisateur_id=utilisateur_id,
        action="telechargement",
        document_id=document.id,
        details=f"Téléchargement de '{document.nom_original}'"
    )

    return document.chemin_fichier


    
       